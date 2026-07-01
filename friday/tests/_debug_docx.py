#!/usr/bin/env python3
"""Minimal test: generate one DOCX from headset research, step by step with full debug."""

import sys, os, asyncio, json, subprocess, textwrap
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))
from dotenv import load_dotenv
load_dotenv(str(Path(__file__).parent.parent / ".env"))
from friday.nim_client import InferenceClient
from friday.nim_router import resolve_model

OUTPUT_DIR = Path(__file__).parent / "_debug_test"
OUTPUT_DIR.mkdir(exist_ok=True)

RESEARCH_PATH = Path(__file__).parent.parent / "memory" / "raw_research_deep_Find_the_best_bluetooth_headsets_under_2000_available_o.md"
research = RESEARCH_PATH.read_text(encoding="utf-8")[:4000]

async def main():
    client = InferenceClient()
    model = resolve_model("code_gen")
    print(f"Model: {model}")
    print(f"Research sample: {len(research)} chars")
    
    # Step 1: Generate python-docx code
    print("\n--- Step 1: Ask LLM to generate code ---")
    prompt = textwrap.dedent(f"""\
        Write Python code using python-docx that creates a document from this research.

        RESEARCH: {research}

        Document: title "Headset Research", one paragraph summary, one table (5 products x 4 cols).
        Save as "test.docx" in current dir.
        ONLY Python code, no markdown.
    """)
    
    resp = await client.chat(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=4000,
    )
    code = resp.content.strip()
    print(f"Response length: {len(code)} chars")
    print(f"First 100 chars: {code[:100]}")
    print(f"Last 100 chars: {code[-100:]}")
    
    if not code:
        print("EMPTY RESPONSE - trying again with different prompt")
        resp = await client.chat(
            model=model,
            messages=[{"role": "user", "content": "Write Python code using python-docx. Create a simple document with title 'Test' and save as 'test.docx'. ONLY code."}],
            temperature=0.1,
            max_tokens=2000,
        )
        code = resp.content.strip()
        print(f"Retry response: {len(code)} chars")
    
    # Extract code
    if "```" in code:
        lines = code.splitlines()
        in_block = False; code_lines = []
        for line in lines:
            if line.startswith("```"): in_block = not in_block; continue
            if in_block: code_lines.append(line)
        code = "\n".join(code_lines) if code_lines else code
    
    fp = OUTPUT_DIR / "gen_docx.py"
    fp.write_text(code, encoding="utf-8")
    print(f"Wrote {len(code)} chars to {fp}")
    
    # Step 2: Run the code
    print("\n--- Step 2: Execute generated code ---")
    result = subprocess.run(
        [sys.executable, str(fp)],
        capture_output=True, text=True, timeout=30,
        cwd=str(OUTPUT_DIR),
    )
    print(f"Exit code: {result.returncode}")
    print(f"stdout: {result.stdout[:500]}")
    print(f"stderr: {result.stderr[:500]}")
    
    # Step 3: Check for output files
    print("\n--- Step 3: Check for output files ---")
    files = list(OUTPUT_DIR.glob("*"))
    for f in files:
        size = f.stat().st_size
        print(f"  {f.name}: {size} bytes")
    
    docx_files = [f for f in files if f.suffix == ".docx"]
    if docx_files:
        from docx import Document
        doc = Document(str(docx_files[0]))
        print(f"DOCX OK: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    else:
        print("NO DOCX FILE CREATED")
        # Check what the code actually did
        print("\n--- Generated code preview ---")
        print(code[:1000])

if __name__ == "__main__":
    asyncio.run(main())
