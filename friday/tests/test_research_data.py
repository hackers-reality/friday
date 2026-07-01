#!/usr/bin/env python3
"""Comprehensive end-to-end test using REAL Bluetooth headset research data to generate PPTX, DOCX, PDF, XLSX."""

import sys, os, json, asyncio, tempfile, textwrap, subprocess, traceback
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from dotenv import load_dotenv
load_dotenv(str(Path(__file__).parent.parent / ".env"))

from friday.nim_client import InferenceClient
from friday.nim_router import resolve_model

PASS = 0
FAIL = 0
RESULTS = []
OUTPUT_DIR = Path(__file__).parent / "_research_test_outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

# Read the Bluetooth headset research file (first 5000 chars as context)
RESEARCH_PATH = Path(__file__).parent.parent / "memory" / "raw_research_deep_Find_the_best_bluetooth_headsets_under_2000_available_o.md"
RESEARCH_DATA = RESEARCH_PATH.read_text(encoding="utf-8")
RESEARCH_SAMPLE = RESEARCH_DATA[:4000] + "\n\n[--- RESEARCH CONTINUES FOR " + str(len(RESEARCH_DATA)) + " TOTAL CHARS ---]"

def log(name, passed, detail=""):
    global PASS, FAIL
    if passed: PASS += 1
    else: FAIL += 1
    RESULTS.append(("PASS" if passed else "FAIL", name, detail))
    print(f"  [{'PASS' if passed else 'FAIL'}] {name}")
    if detail: print(f"         {detail[:120]}")

async def llm_gen(prompt, model=None, temp=0.1, max_tokens=6000):
    client = InferenceClient()
    if model is None: model = resolve_model("code_gen")
    resp = await client.chat(model=model, messages=[{"role": "user", "content": prompt}], temperature=temp, max_tokens=max_tokens)
    code = resp.content.strip()
    if not code:
        # Retry once with higher temperature
        resp = await client.chat(model=model, messages=[{"role": "user", "content": prompt + "\n\nIMPORTANT: Output the Python code now. Do NOT return empty response."}], temperature=0.3, max_tokens=max_tokens)
        code = resp.content.strip()
    return code

def extract_code(text):
    if "```" not in text: return text
    lines = text.splitlines()
    in_block = False; code_lines = []
    for line in lines:
        if line.startswith("```"):
            in_block = not in_block; continue
        if in_block: code_lines.append(line)
    return "\n".join(code_lines) if code_lines else text

async def test_research_pptx():
    """Generate PPTX presentation from Bluetooth headset research."""
    print("\n=== REAL RESEARCH TEST: PPTX Presentation ===")
    prompt = textwrap.dedent(f"""\
        Using the research data below about Bluetooth headsets under ₹2000, write Python code
        using python-pptx that creates a 6-slide presentation:

        RESEARCH DATA:
        {RESEARCH_SAMPLE}

        Slide 1: Title "Best Bluetooth Headsets Under ₹2000 (2025)" subtitle "Comprehensive Market Research"
        Slide 2: "Top 5 Picks" with a table showing: Brand, Model, Price, Battery Life, Rating
        Slide 3: "Key Features Comparison" - bullet list of important features
        Slide 4: "Price vs Performance" - chart overview
        Slide 5: "Buying Guide" - what to look for when buying budget Bluetooth headsets
        Slide 6: "Recommendations" - final verdict and best buys

        Design: font floor 14pt, professional colors, consistent layouts.
        Save as "headset_research.pptx" in current directory.
        Output ONLY Python code, no markdown wrapper.
    """)
    try:
        code = await llm_gen(prompt, temp=0.3, max_tokens=6000)
        code = extract_code(code)
        fp = OUTPUT_DIR / "gen_pptx.py"
        fp.write_text(code, encoding="utf-8")
        result = subprocess.run([sys.executable, str(fp)], capture_output=True, text=True, timeout=60, cwd=str(OUTPUT_DIR))
        pptx_files = list(OUTPUT_DIR.glob("*.pptx"))
        if result.returncode == 0 and pptx_files:
            from pptx import Presentation
            prs = Presentation(str(pptx_files[0]))
            log("Research PPTX", True, f"{len(prs.slides)} slides, {pptx_files[0].name}")
        else:
            log("Research PPTX", False, f"exit={result.returncode}, err={result.stderr[:200]}")
    except Exception as e:
        log("Research PPTX", False, str(e)[:200])

async def test_research_docx():
    """Generate DOCX report from Bluetooth headset research."""
    print("\n=== REAL RESEARCH TEST: DOCX Report ===")
    prompt = textwrap.dedent(f"""\
        Using the research data below about Bluetooth headsets under ₹2000, write Python code
        using python-docx that creates a professional report:

        RESEARCH DATA:
        {RESEARCH_SAMPLE}

        Document structure:
        - Title: "Bluetooth Headset Market Analysis — Under ₹2000"
        - Section 1: Executive Summary (paragraph summarizing the research)
        - Section 2: Top Products Table (5 rows × 5 columns: Brand, Model, Price, Battery, Rating)
        - Section 3: Feature Comparison (bullet list of key features found in the research)
        - Section 4: Buying Recommendations (numbered list of tips)
        - Section 5: Conclusion

        Design: A4, 1-inch margins, page numbers in footer, professional fonts.
        Save as "headset_research.docx" in current directory.
        Output ONLY Python code, no markdown wrapper.
    """)
    try:
        code = await llm_gen(prompt, temp=0.3, max_tokens=6000)
        code = extract_code(code)
        fp = OUTPUT_DIR / "gen_docx.py"
        fp.write_text(code, encoding="utf-8")
        result = subprocess.run([sys.executable, str(fp)], capture_output=True, text=True, timeout=60, cwd=str(OUTPUT_DIR))
        docx_files = list(OUTPUT_DIR.glob("*.docx"))
        if result.returncode == 0 and docx_files:
            from docx import Document
            doc = Document(str(docx_files[0]))
            log("Research DOCX", True, f"{len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        else:
            log("Research DOCX", False, f"exit={result.returncode}, err={result.stderr[:200]}")
    except Exception as e:
        log("Research DOCX", False, str(e)[:200])

async def test_research_pdf():
    """Generate PDF report from Bluetooth headset research using reportlab."""
    print("\n=== REAL RESEARCH TEST: PDF Report ===")
    prompt = textwrap.dedent(f"""\
        Using the research data below about Bluetooth headsets under ₹2000, write Python code
        using reportlab that creates a 3-page PDF report:

        RESEARCH DATA:
        {RESEARCH_SAMPLE}

        Page 1: Title page "Bluetooth Headset Market Report 2025" with subtitle and date
        Page 2: Key findings with a table (5 products, columns: Brand, Model, Price, Rating)
        Page 3: Recommendations and buying guide with bullet points

        Include: page numbers in footer, professional styling, proper margins.
        Save as "headset_research.pdf" in current directory.
        Output ONLY Python code, no markdown wrapper.
    """)
    try:
        code = await llm_gen(prompt, temp=0.3, max_tokens=6000)
        code = extract_code(code)
        fp = OUTPUT_DIR / "gen_pdf.py"
        fp.write_text(code, encoding="utf-8")
        result = subprocess.run([sys.executable, str(fp)], capture_output=True, text=True, timeout=60, cwd=str(OUTPUT_DIR))
        pdf_files = list(OUTPUT_DIR.glob("*.pdf"))
        if result.returncode == 0 and pdf_files:
            import PyPDF2
            with open(str(pdf_files[0]), "rb") as f:
                reader = PyPDF2.PdfReader(f)
                log("Research PDF", True, f"{len(reader.pages)} pages, {pdf_files[0].name}")
        else:
            log("Research PDF", False, f"exit={result.returncode}, err={result.stderr[:200]}")
    except Exception as e:
        log("Research PDF", False, str(e)[:200])

async def test_research_xlsx():
    """Generate XLSX spreadsheet from Bluetooth headset research."""
    print("\n=== REAL RESEARCH TEST: XLSX Spreadsheet ===")
    prompt = textwrap.dedent(f"""\
        Using the research data below about Bluetooth headsets under ₹2000, write Python code
        using openpyxl that creates a spreadsheet with product comparison:

        RESEARCH DATA:
        {RESEARCH_SAMPLE}

        Sheet 1 "Product Comparison":
        - Header row: Product, Brand, Price, Battery Life, Rating, Bluetooth Version, Warranty, Buy Link
        - 8-10 data rows with real product data from the research
        - Bold headers with blue fill, alternating row colors
        - Auto-adjusted column widths

        Sheet 2 "Price Analysis":
        - Header: Product, Price, Value Rating
        - Sort by price ascending

        Design: freeze panes, conditional formatting for price ranges, data validation.
        Save as "headset_comparison.xlsx" in current directory.
        Output ONLY Python code, no markdown wrapper.
    """)
    try:
        code = await llm_gen(prompt, temp=0.3, max_tokens=6000)
        code = extract_code(code)
        fp = OUTPUT_DIR / "gen_xlsx.py"
        fp.write_text(code, encoding="utf-8")
        result = subprocess.run([sys.executable, str(fp)], capture_output=True, text=True, timeout=60, cwd=str(OUTPUT_DIR))
        xlsx_files = list(OUTPUT_DIR.glob("*.xlsx"))
        if result.returncode == 0 and xlsx_files:
            from openpyxl import load_workbook
            wb = load_workbook(str(xlsx_files[0]))
            log("Research XLSX", True, f"sheets={wb.sheetnames}, rows={wb.active.max_row}")
        else:
            log("Research XLSX", False, f"exit={result.returncode}, err={result.stderr[:200]}")
    except Exception as e:
        log("Research XLSX", False, str(e)[:200])

async def test_research_svg_animated():
    """Generate animated FRIDAY logo SVG."""
    print("\n=== REAL RESEARCH TEST: Animated SVG Logo for FRIDAY ===")
    prompt = textwrap.dedent("""\
        Generate Python code that writes an animated SVG logo file for "FRIDAY AI" to disk.

        The SVG must be an animated logo with:
        - The text "FRIDAY" in bold modern font, with a gradient fill
        - A circular ring around the text that rotates continuously (CSS animation)
        - Small orbiting dots that pulse (CSS keyframe animations)
        - The subtitle "AI ASSISTANT" below in smaller text
        - Color scheme: dark background (#0a0a2e), cyan/electric blue accents (#00d4ff, #0066ff)
        - viewBox responsive sizing
        - Pure CSS animations (no JavaScript)
        - Professional, modern tech aesthetic

        Save as "friday_animated_logo.svg" in current directory.
        Output ONLY Python code, no markdown wrapper.
    """)
    try:
        code = await llm_gen(prompt, temp=0.3, max_tokens=4000)
        code = extract_code(code)
        fp = OUTPUT_DIR / "gen_svg.py"
        fp.write_text(code, encoding="utf-8")
        result = subprocess.run([sys.executable, str(fp)], capture_output=True, text=True, timeout=60, cwd=str(OUTPUT_DIR))
        svg_files = list(OUTPUT_DIR.glob("*.svg"))
        if result.returncode == 0 and svg_files:
            content = svg_files[0].read_text(encoding="utf-8")
            has_svg = "<svg" in content and "</svg>" in content
            has_animation = "@keyframes" in content or "<animate" in content or "animation" in content
            log("Animated SVG Logo", has_svg, f"animated={has_animation}, {len(content)} chars")
        else:
            log("Animated SVG Logo", False, f"exit={result.returncode}, err={result.stderr[:200]}")
    except Exception as e:
        log("Animated SVG Logo", False, str(e)[:200])

async def test_forge_complex_task():
    """Test Forge with a complex multi-file coding task using research data."""
    print("\n=== REAL RESEARCH TEST: Forge Multi-File Project ===")
    prompt = textwrap.dedent(f"""\
        Generate Python code for a complete data analysis project using the headset research data.
        The project should have two files:

        # FILE: analysis.py
        - Define 5-6 Bluetooth headset products with real data (name, brand, price, battery, rating)
        - Functions to sort by price, filter by rating, calculate average price
        - Print formatted analysis results

        # FILE: test_analysis.py
        - Unit tests for all functions in analysis.py using pytest
        - Test edge cases (empty list, single item, etc.)

        RESEARCH DATA SAMPLE:
        {RESEARCH_SAMPLE[:3000]}

        Output code with "# FILE: filename.py" separators.
    """)
    try:
        code = await llm_gen(prompt, temp=0.3, max_tokens=6000)
        code = extract_code(code)
        # Parse multi-file output
        files_written = []
        current_file = None; current_content = []
        for line in code.splitlines():
            if line.startswith("# FILE:") or line.startswith("# file:"):
                if current_file and current_content:
                    (OUTPUT_DIR / current_file).write_text("\n".join(current_content), encoding="utf-8")
                    files_written.append(current_file)
                current_file = line.split(":")[-1].strip().strip('"').strip("'")
                current_content = []
            else:
                current_content.append(line)
        if current_file and current_content:
            (OUTPUT_DIR / current_file).write_text("\n".join(current_content), encoding="utf-8")
            files_written.append(current_file)

        if not files_written:
            (OUTPUT_DIR / "analysis.py").write_text(code, encoding="utf-8")
            files_written.append("analysis.py")

        # Check syntax of all files
        ok = True
        for fname in files_written:
            r = subprocess.run([sys.executable, "-c", f"import ast; ast.parse(open(r'{OUTPUT_DIR / fname}').read())"],
                             capture_output=True, text=True, timeout=15)
            if r.returncode != 0:
                ok = False; log("Forge Multi-File", False, f"{fname} syntax error: {r.stderr[:150]}"); break
        if ok:
            log("Forge Multi-File", True, f"{len(files_written)} files, syntax OK")
    except Exception as e:
        log("Forge Multi-File", False, str(e)[:200])

async def main():
    print("=" * 60)
    print("FRIDAY COMPREHENSIVE RESEARCH-DRIVEN TEST SUITE")
    print(f"Research file: {RESEARCH_PATH.name} ({len(RESEARCH_DATA)} chars)")
    print(f"Using model: {resolve_model('code_gen')}")
    print("=" * 60)
    print()

    await test_research_pptx()
    await test_research_docx()
    await test_research_pdf()
    await test_research_xlsx()
    await test_research_svg_animated()
    await test_forge_complex_task()

    print(f"\n{'='*60}")
    print(f"  FINAL RESULTS: {PASS} passed, {FAIL} failed out of {PASS+FAIL}")
    print(f"{'='*60}")
    print(f"  Output files in: {OUTPUT_DIR}")
    for s, name, d in RESULTS:
        print(f"  [{s}] {name:<40} | {d[:100]}")

if __name__ == "__main__":
    asyncio.run(main())
