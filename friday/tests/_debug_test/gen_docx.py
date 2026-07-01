from docx import Document
from docx.shared import Pt, Inches, Cm

# Create a new Document object
doc = Document()

# Set document default font (optional)
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ----- Title -----
doc.add_heading('Headset Research', level=1)

# ----- Summary Paragraph -----
summary = (
    "The research aimed to find the best Bluetooth headsets under 2000 rupees in India across "
    "Amazon, Flipkart, and official manufacturer sites. After scraping 46 pages and collecting "
    "284 links, a list of top products was compiled based on customer reviews and deals. "
    "The top 5 products are presented in the table below."
)
doc.add_paragraph(summary)

# ----- Table (5 products x 4 columns) -----
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Shading Accent 1'  # built-in style

# Header row
header_cells = table.rows[0].cells
headers = ['Product', 'Price (₹)', 'Platform', 'Rating']
for i, header in enumerate(headers):
    header_cells[i].text = header
    # Make header text bold
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows (5 products)
products = [
    ['boAt Rockerz 480',        '1,499', 'Amazon / Flipkart', '4.2 / 5'],
    ['Realme Buds Air 5',       '1,999', 'Flipkart',          '4.4 / 5'],
    ['Noise Buds N1',           '1,299', 'Amazon',            '4.1 / 5'],
    ['JBL Tune 510BT',          '1,949', 'Amazon',            '4.0 / 5'],
    ['Snapods View TWS',        '1,799', 'Snapup Official',   '4.3 / 5'],
]

for product_data in products:
    row_cells = table.add_row().cells
    for idx, value in enumerate(product_data):
        row_cells[idx].text = value

# Save the document
doc.save('test.docx')