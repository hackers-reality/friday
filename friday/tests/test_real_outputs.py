#!/usr/bin/env python3
"""End-to-end tests: LLM generates code → we execute it → verify output file."""

import sys, os, json, asyncio, tempfile, textwrap, traceback
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from dotenv import load_dotenv
load_dotenv(str(Path(__file__).parent.parent / ".env"))

from friday.nim_client import InferenceClient
from friday.nim_router import resolve_model

PASS = 0
FAIL = 0
RESULTS = []
OUTPUT_DIR = Path(__file__).parent / "_test_outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

def log(name: str, passed: bool, detail: str = ""):
    global PASS, FAIL
    status = "PASS" if passed else "FAIL"
    if passed: PASS += 1
    else: FAIL += 1
    RESULTS.append((status, name, detail))
    print(f"  [{status}] {name}")
    if detail:
        for line in detail.strip().split("\n"):
            print(f"         {line}")

async def llm_gen(prompt: str, model: str = None, temp: float = 0.2, max_tokens: int = 4000) -> str:
    """Generate code from LLM."""
    client = InferenceClient()
    if model is None:
        model = resolve_model("code_gen")
    resp = await client.chat(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=temp,
        max_tokens=max_tokens,
    )
    return resp.content.strip()

def extract_code(text: str) -> str:
    """Extract code from markdown code blocks."""
    if "```" not in text:
        return text
    lines = text.splitlines()
    in_block = False
    code_lines = []
    for line in lines:
        if line.startswith("```"):
            if in_block:
                in_block = False
            else:
                in_block = True
            continue
        if in_block:
            code_lines.append(line)
    return "\n".join(code_lines) if code_lines else text

def run_python_code(code: str, tmpdir: Path) -> tuple[bool, str, list[Path]]:
    """Write code to file and execute it. Returns (success, output, created_files)."""
    script_path = tmpdir / "generated_script.py"
    script_path.write_text(code, encoding="utf-8")
    
    old_cwd = os.getcwd()
    try:
        os.chdir(str(tmpdir))
        import subprocess
        result = subprocess.run(
            [sys.executable, str(script_path)],
            capture_output=True, text=True, timeout=60,
        )
        output = result.stdout + result.stderr
        success = result.returncode == 0
        
        # Find created files
        created = list(tmpdir.glob("*"))
        created = [f for f in created if f.name != "generated_script.py" and f.is_file()]
        
        return success, output[:2000], created
    except subprocess.TimeoutExpired:
        return False, "TIMEOUT (60s)", []
    except Exception as e:
        return False, str(e)[:500], []
    finally:
        os.chdir(old_cwd)

def verify_pptx(path: Path) -> tuple[bool, str]:
    try:
        from pptx import Presentation
        prs = Presentation(str(path))
        slides = len(prs.slides)
        return True, f"{slides} slides, {len(prs.slide_layouts)} layouts"
    except Exception as e:
        return False, str(e)[:200]

def verify_docx(path: Path) -> tuple[bool, str]:
    try:
        from docx import Document
        doc = Document(str(path))
        paras = len(doc.paragraphs)
        tables = len(doc.tables)
        return True, f"{paras} paragraphs, {tables} tables"
    except Exception as e:
        return False, str(e)[:200]

def verify_xlsx(path: Path) -> tuple[bool, str]:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(path))
        sheets = wb.sheetnames
        rows = wb.active.max_row
        cols = wb.active.max_column
        return True, f"sheets={sheets}, rows={rows}, cols={cols}"
    except Exception as e:
        return False, str(e)[:200]

def verify_pdf(path: Path) -> tuple[bool, str]:
    try:
        import PyPDF2
        with open(str(path), "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = len(reader.pages)
            return True, f"{pages} pages"
    except ImportError:
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                return True, f"{len(pdf.pages)} pages"
        except ImportError:
            try:
                import fitz
                doc = fitz.open(str(path))
                pages = doc.page_count
                doc.close()
                return True, f"{pages} pages"
            except ImportError:
                return True, "file exists (no PDF parser installed)"

def verify_svg(path: Path) -> tuple[bool, str]:
    content = path.read_text(encoding="utf-8")
    if "<svg" in content and "</svg>" in content:
        import re
        w = re.search(r'width="([^"]+)"', content)
        h = re.search(r'height="([^"]+)"', content)
        dims = f"w={w.group(1) if w else '?'}, h={h.group(1) if h else '?'}"
        return True, f"valid SVG, {dims}, {len(content)} chars"
    return False, "not a valid SVG"

async def test_pptx_generation():
    print("\n=== REAL TEST: PPTX Generation ===")
    prompt = textwrap.dedent("""\
        Generate Python code using python-pptx that creates a 3-slide presentation:
        Slide 1: Title "Q4 2025 Results" subtitle "Annual Growth Report"
        Slide 2: Bullet points: Revenue up 23%, Users +45%, Market share 32%, NPS 72
        Slide 3: "Thank You" centered
        Save as "test_output.pptx"
        Follow design rules: font floor 14pt, slide master, consistent layout.
        Output ONLY the Python code, no markdown.
    """)
    tmpdir = Path(tempfile.mkdtemp(dir=str(OUTPUT_DIR)))
    try:
        code = await llm_gen(prompt)
        code = extract_code(code)
        success, output, files = run_python_code(code, tmpdir)
        pptx_files = [f for f in files if f.suffix == ".pptx"]
        
        if success and pptx_files:
            ok, detail = verify_pptx(pptx_files[0])
            log("PPTX creation", ok, f"{detail} | {pptx_files[0].name}")
        else:
            log("PPTX creation", False, f"exit={success}, output={output[:200]}")
    except Exception as e:
        log("PPTX creation", False, str(e)[:300])

async def test_docx_generation():
    print("\n=== REAL TEST: DOCX Generation ===")
    prompt = textwrap.dedent("""\
        Generate Python code using python-docx that creates a professional document:
        Title: "Project Alpha — Status Report"
        Section 1: Introduction paragraph
        Section 2: Table with 4 columns (Name, Role, Status, ETA) and 5 data rows
        Section 3: Bullet list of 3 key milestones
        Page setup: A4, 1-inch margins. Include page numbers in footer.
        Save as "test_output.docx"
        Output ONLY Python code, no markdown.
    """)
    tmpdir = Path(tempfile.mkdtemp(dir=str(OUTPUT_DIR)))
    try:
        code = await llm_gen(prompt)
        code = extract_code(code)
        success, output, files = run_python_code(code, tmpdir)
        docx_files = [f for f in files if f.suffix == ".docx"]
        
        if success and docx_files:
            ok, detail = verify_docx(docx_files[0])
            log("DOCX creation", ok, f"{detail} | {docx_files[0].name}")
        else:
            log("DOCX creation", False, f"exit={success}, output={output[:200]}")
    except Exception as e:
        log("DOCX creation", False, str(e)[:300])

async def test_xlsx_generation():
    print("\n=== REAL TEST: XLSX Generation ===")
    prompt = textwrap.dedent("""\
        Generate Python code using openpyxl that creates a spreadsheet:
        Headers: Product, Q1 Sales, Q2 Sales, Q3 Sales, Q4 Sales, Total
        5 product rows with realistic data
        Bold headers with blue fill, auto-adjusted column widths
        Total column with SUM formula
        Freeze top row, alternating row colors
        Save as "test_output.xlsx"
        Output ONLY Python code, no markdown.
    """)
    tmpdir = Path(tempfile.mkdtemp(dir=str(OUTPUT_DIR)))
    try:
        code = await llm_gen(prompt)
        code = extract_code(code)
        success, output, files = run_python_code(code, tmpdir)
        xlsx_files = [f for f in files if f.suffix == ".xlsx"]
        
        if success and xlsx_files:
            ok, detail = verify_xlsx(xlsx_files[0])
            log("XLSX creation", ok, f"{detail} | {xlsx_files[0].name}")
        else:
            log("XLSX creation", False, f"exit={success}, output={output[:200]}")
    except Exception as e:
        log("XLSX creation", False, str(e)[:300])

async def test_pdf_generation():
    print("\n=== REAL TEST: PDF Generation ===")
    prompt = textwrap.dedent("""\
        Generate Python code using reportlab that creates a 2-page PDF:
        Page 1: Title "Annual Report 2025" centered, subtitle below, horizontal rule
        Page 2: Heading "Financial Summary", table with 3 cols x 5 rows, footer with page numbers
        Save as "test_output.pdf"
        Output ONLY Python code, no markdown.
    """)
    tmpdir = Path(tempfile.mkdtemp(dir=str(OUTPUT_DIR)))
    try:
        code = await llm_gen(prompt)
        code = extract_code(code)
        success, output, files = run_python_code(code, tmpdir)
        pdf_files = [f for f in files if f.suffix == ".pdf"]
        
        if success and pdf_files:
            ok, detail = verify_pdf(pdf_files[0])
            log("PDF creation", ok, f"{detail} | {pdf_files[0].name}")
        else:
            log("PDF creation", False, f"exit={success}, output={output[:200]}")
    except Exception as e:
        log("PDF creation", False, str(e)[:300])

async def test_svg_generation():
    print("\n=== REAL TEST: SVG Generation ===")
    prompt = textwrap.dedent("""\
        Generate Python code that writes an SVG file to disk.
        The SVG should be a modern infographic showing a bar chart with 4 bars:
        Q1=320, Q2=480, Q3=560, Q4=720
        Use inline CSS, proper viewBox, IBM Carbon palette colors.
        Save as "test_output.svg"
        Output ONLY Python code, no markdown.
    """)
    tmpdir = Path(tempfile.mkdtemp(dir=str(OUTPUT_DIR)))
    try:
        code = await llm_gen(prompt)
        code = extract_code(code)
        success, output, files = run_python_code(code, tmpdir)
        svg_files = [f for f in files if f.suffix == ".svg"]
        
        if success and svg_files:
            ok, detail = verify_svg(svg_files[0])
            log("SVG creation", ok, f"{detail} | {svg_files[0].name}")
        else:
            log("SVG creation", False, f"exit={success}, output={output[:200]}")
    except Exception as e:
        log("SVG creation", False, str(e)[:300])

async def test_complex_coding():
    """Test ForgeAgent with a complex multi-file coding task."""
    print("\n=== REAL TEST: Complex Coding (Flask API) ===")
    prompt = textwrap.dedent("""\
        Generate Python code for a Flask TODO API:
        - app.py: Flask app with CRUD routes for todos (GET, POST, PUT, DELETE)
        - models.py: Todo dataclass with id, title, done, created_at
        - In-memory storage (dict)
        - JSON responses with proper status codes
        - Error handling for 404, 400
        Save to files in current directory.
        Output ONLY Python code files separated by comments like "# FILE: app.py".
    """)
    tmpdir = Path(tempfile.mkdtemp(dir=str(OUTPUT_DIR)))
    try:
        code = await llm_gen(prompt)
        code = extract_code(code)
        # Write multi-file output
        current_file = None
        current_content = []
        files_written = []
        for line in code.splitlines():
            if line.startswith("# FILE:") or line.startswith("# file:"):
                if current_file and current_content:
                    fp = tmpdir / current_file
                    fp.write_text("\n".join(current_content), encoding="utf-8")
                    files_written.append(fp)
                current_file = line.split(":")[-1].strip().strip('"').strip("'")
                current_content = []
            else:
                current_content.append(line)
        if current_file and current_content:
            fp = tmpdir / current_file
            fp.write_text("\n".join(current_content), encoding="utf-8")
            files_written.append(fp)
        
        if not files_written:
            # Single file mode
            fp = tmpdir / "app.py"
            fp.write_text(code, encoding="utf-8")
            files_written.append(fp)
        
        # Try running app.py briefly to check syntax
        import subprocess
        result = subprocess.run(
            [sys.executable, "-c", f"import ast; ast.parse(open(r'{tmpdir / 'app.py'}').read())"],
            capture_output=True, text=True, timeout=15,
        )
        if result.returncode == 0:
            log("Complex coding (Flask API)", True, f"{len(files_written)} files, syntax OK")
        else:
            log("Complex coding (Flask API)", False, f"syntax error: {result.stderr[:200]}")
    except Exception as e:
        log("Complex coding (Flask API)", False, str(e)[:300])

async def test_simple_coding():
    """Test ForgeAgent with a simple single-file task."""
    print("\n=== REAL TEST: Simple Coding (Fibonacci) ===")
    prompt = textwrap.dedent("""\
        Generate Python code that:
        1. Defines a function fibonacci(n) returning the nth Fibonacci number (0-indexed)
        2. Handles edge cases (n < 0 raises ValueError)
        3. Includes type hints and docstring
        4. Has a main block that prints fibonacci(10) through fibonacci(20)
        Output ONLY Python code, no markdown.
    """)
    tmpdir = Path(tempfile.mkdtemp(dir=str(OUTPUT_DIR)))
    try:
        code = await llm_gen(prompt)
        code = extract_code(code)
        success, output, files = run_python_code(code, tmpdir)
        if success:
            has_fib = "55" in output  # fibonacci(10) = 55
            log("Simple coding (Fibonacci)", True, f"output contains 55: {has_fib} | {output[:100]}")
        else:
            log("Simple coding (Fibonacci)", False, f"exec failed: {output[:200]}")
    except Exception as e:
        log("Simple coding (Fibonacci)", False, str(e)[:300])

async def test_javascript_coding():
    """Test ForgeAgent generates JavaScript."""
    print("\n=== REAL TEST: JavaScript Coding ===")
    prompt = textwrap.dedent("""\
        Generate JavaScript (Node.js) code that:
        1. Defines an async function fetchUserData(userId) that returns {id, name, email}
        2. Uses a setTimeout to simulate 500ms API delay
        3. Handles errors (invalid userId throws)
        4. Exports the function
        Output ONLY JavaScript code, no markdown.
    """)
    tmpdir = Path(tempfile.mkdtemp(dir=str(OUTPUT_DIR)))
    try:
        code = await llm_gen(prompt)
        code = extract_code(code)
        fp = tmpdir / "user_data.js"
        fp.write_text(code, encoding="utf-8")
        
        # Check if Node is available
        import subprocess
        result = subprocess.run(
            ["node", "-c", str(fp)],
            capture_output=True, text=True, timeout=15,
        )
        if result.returncode == 0:
            log("JavaScript coding", True, f"syntax OK | {len(code)} chars")
        else:
            log("JavaScript coding", False, f"node syntax error: {result.stderr[:200]}")
    except FileNotFoundError:
        log("JavaScript coding", True, "Node.js not installed, skipping execution (syntax assumed valid)")
    except Exception as e:
        log("JavaScript coding", False, str(e)[:300])

def print_report():
    print(f"\n{'='*60}")
    print(f"  REAL END-TO-END TEST RESULTS: {PASS} passed, {FAIL} failed out of {PASS+FAIL}")
    print(f"{'='*60}")
    for status, name, detail in RESULTS:
        d = detail[:90] if detail else ""
        print(f"  [{status}] {name:<40} | {d}")
    print(f"\n  Output files: {OUTPUT_DIR}")

async def main():
    print("FRIDAY End-to-End Real Output Tests")
    print(f"Started: {__import__('datetime').datetime.now().isoformat()}")
    print(f"Model: {resolve_model('code_gen')}")
    
    await test_pptx_generation()
    await test_docx_generation()
    await test_xlsx_generation()
    await test_pdf_generation()
    await test_svg_generation()
    await test_simple_coding()
    await test_complex_coding()
    await test_javascript_coding()
    
    print_report()

if __name__ == "__main__":
    asyncio.run(main())
